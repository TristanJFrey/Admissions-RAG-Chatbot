"""
convert_docx_to_json.py

Converts a structured DOCX document (with Headings, key-value pairs, bullet lists, and tables)
into a hierarchical JSON file. Designed for converting structured reports into machine-usable data.

Usage:
    python convert_docx_to_json.py input.docx output.json
"""

import sys
import json
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from typing import Dict, Any, List

def is_heading(p: Paragraph) -> bool:
    name = p.style.name if p.style is not None else ""
    return name.startswith("Heading")

def heading_level(p: Paragraph) -> int:
    try:
        return int(p.style.name.split()[-1])
    except Exception:
        return 9

def is_list_bullet(p: Paragraph) -> bool:
    name = p.style.name if p.style is not None else ""
    return ("List Bullet" in name) or ("List Number" in name)

def parse_kv_from_paragraph(p: Paragraph):
    text = p.text.strip()
    if not text:
        return None
    if p.runs and any(r.bold for r in p.runs):
        bold = ""
        for r in p.runs:
            if r.bold:
                bold += r.text
            else:
                break
        candidate = bold.strip().rstrip(":")
        if candidate:
            if ":" in text:
                key, val = text.split(":", 1)
                return key.strip(), val.strip()
            else:
                val = text[len(bold):].strip(" :")
                if val:
                    return candidate, val
    if ":" in text:
        key, val = text.split(":", 1)
        if key.strip() and val.strip():
            return key.strip(), val.strip()
    return None

def table_to_list(tbl: Table):
    rows = tbl.rows
    if not rows:
        return []
    headers = [c.text.strip() or f"col_{i+1}" for i, c in enumerate(rows[0].cells)]
    data = []
    for r in rows[1:]:
        obj = {h: c.text.strip() for h, c in zip(headers, r.cells)}
        data.append(obj)
    return data

def convert_docx_to_json(input_path: str, output_path: str):
    doc = Document(input_path)
    root: Dict[str, Any] = {"title": None, "content": []}
    stack = [(0, root)]
    pending_list: List[str] = None

    def push_section(title: str, level: int) -> Dict[str, Any]:
        node = {"title": title, "level": level, "content": []}
        while stack and stack[-1][0] >= level:
            stack.pop()
        stack[-1][1]["content"].append(node)
        stack.append((level, node))
        return node

    def finalize_list_if_any(target_node):
        nonlocal pending_list
        if pending_list is not None:
            target_node["content"].append({"type": "list", "items": pending_list})
            pending_list = None

    if doc.paragraphs:
        first = doc.paragraphs[0]
        if first.style and first.style.name == "Title":
            root["title"] = first.text.strip()

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text:
                continue
            if is_heading(p):
                finalize_list_if_any(stack[-1][1])
                lvl = heading_level(p)
                push_section(text, lvl)
                continue
            kv = parse_kv_from_paragraph(p)
            if kv:
                finalize_list_if_any(stack[-1][1])
                k, v = kv
                stack[-1][1]["content"].append({"type": "kv", "key": k, "value": v})
                continue
            if is_list_bullet(p):
                if pending_list is None:
                    pending_list = []
                pending_list.append(text)
                continue
            finalize_list_if_any(stack[-1][1])
            stack[-1][1]["content"].append({"type": "paragraph", "text": text})
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            finalize_list_if_any(stack[-1][1])
            stack[-1][1]["content"].append({"type": "table", "rows": table_to_list(tbl)})

    finalize_list_if_any(stack[-1][1])

    def clean(node: Dict[str, Any]):
        node2 = {}
        for k, v in node.items():
            if k == "level":
                continue
            if k == "content":
                node2[k] = []
                for item in v:
                    if isinstance(item, dict) and "title" in item and "content" in item:
                        node2[k].append(clean(item))
                    else:
                        node2[k].append(item)
            else:
                node2[k] = v
        return node2

    result = clean(root)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"Converted {input_path} -> {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert_docx_to_json.py input.docx output.json")
        sys.exit(1)
    convert_docx_to_json(sys.argv[1], sys.argv[2])
